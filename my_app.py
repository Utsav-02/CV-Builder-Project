from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)


document = Document()


# Profle Picture (You can Adjust Inches as per your need)

document.add_picture(
    'image.jpg' , 
    width=Inches(2.0))

# Name, Phone no. , email detials

name = input('Enter your name - ')
speak('Hello' + name + 'how are you tody?')

speak('Enter your phone number - ')
phone_number = str(input('Enter your phone number - '))
email = input('Enter Your Email - ')

document.add_paragraph(
name + ' | ' + phone_number + ' | ' + email )

# about me

document.add_heading("ABOUT ME")
about_me = input('Tell me about yourself? ')
document.add_paragraph(about_me)

# Work Experiences
document.add_heading("WORK EXPERIENCE")
P = document.add_paragraph()


company = input('Enter Company Name - ')
from_date = input('From Date - ')
to_date = input('To date - ')


P.add_run(company + ' ' ).bold = True
P.add_run(from_date + ' - ' + to_date + '\n' ).italic = True


experience_details = input( \
'Describe your experience at ' + company  )

P.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
      'Do you have more exeperiences?  Yes or No - ')
    if has_more_experiences.lower() == 'yes':
        P = document.add_paragraph()

        company = input('Enter Company Name - ')
        from_date = input('From Date - ')
        to_date = input('To date - ')
        
        P.add_run(company + ' ' ).bold = True
        P.add_run(from_date + ' - ' + to_date + '\n' ).italic = True
        
        experience_details = input('' \
            'Describe your experience at ' + company)
        P.add_run(experience_details)


    else:
        break



# SKILLS
document.add_heading('SKILLS')
Skill = input('Enter your Skill - ')

P = document.add_paragraph(Skill)
P.style = 'List Bullet'

while True:
    has_more_skills = input('' \
    'Do you have more skills Yes or No - ' )
    
    if has_more_skills.lower() == 'yes':
        Skill = input('Enter your Skill - ')
        
        P = document.add_paragraph(Skill)
        P.style = 'List Bullet' 
    
    else:
        break
        
#footer
section = document.sections[0]
footer = section.footer
P = footer.paragraphs[0]
P.text = 'CV generated using Python'



document.save('cv.docx')
